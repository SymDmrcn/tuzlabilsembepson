import pandas as pd
import streamlit as st
import urllib.request
from io import BytesIO
from docx import Document

# 📌 Streamlit Arayüzü
st.title("Tuzla BİLSEM BEP Hazırlama Uygulaması")

# Kullanıcı giriş alanları
teacher = st.text_input("Öğretmen Adı:")
student = st.text_input("Öğrenci Adı:")

# 📌 GitHub’daki Excel Dosyasını İndir ve Oku
excel_url = "https://raw.githubusercontent.com/SymDmrcn/bepuygulama/main/VERİLER.xlsx"

try:
    # 🔹 Excel dosyasını UTF-8 ile oku (Burası önemli!)
    response = urllib.request.urlopen(excel_url)
    excel_data = BytesIO(response.read())  # Belleğe al
    df = pd.read_excel(excel_data, sheet_name="Sayfa1", skiprows=1, engine="openpyxl")

    # 🔹 Tüm metinleri UTF-8 formatına çevir
    df = df.astype(str).applymap(lambda x: x.encode("utf-8", "ignore").decode("utf-8"))

    # 🔹 Sütun isimlerini düzenle
    df.columns = ["GRUP", "DERS", "KISA VADELİ HEDEFLER", "UZUN VADELİ HEDEFLER", "ÖĞRETİMSEL HEDEFLER"]
    df = df.dropna(how="all")

except Exception as e:
    st.error(f"❌ Excel dosyası yüklenirken hata oluştu:\n\n{e}")
    st.stop()

# 📌 Grup seçimi
groups = sorted(df["GRUP"].dropna().unique())
selected_group = st.selectbox("Grup Seçin:", groups)

# 📌 Seçilen gruba göre ders listesini getir
if selected_group:
    filtered_df = df[df["GRUP"] == selected_group]
    lessons = sorted(filtered_df["DERS"].dropna().unique())
    selected_lesson = st.selectbox("Ders Seçin:", lessons)

    if selected_lesson:
        lesson_data = filtered_df[filtered_df["DERS"] == selected_lesson]
        
        st.subheader("Kısa Vadeli Hedefler")
        short_terms = lesson_data["KISA VADELİ HEDEFLER"].dropna().astype(str).values
        selected_short_terms = st.multiselect("Seçiniz:", short_terms)

        st.subheader("Uzun Vadeli Hedefler")
        long_terms = lesson_data["UZUN VADELİ HEDEFLER"].dropna().astype(str).values
        selected_long_terms = st.multiselect("Seçiniz:", long_terms)

        st.subheader("Öğretimsel Hedefler")
        teaching_terms = lesson_data["ÖĞRETİMSEL HEDEFLER"].dropna().astype(str).values
        selected_teaching_terms = st.multiselect("Seçiniz:", teaching_terms)

        # 📌 Seçimleri Word Belgesine Kaydetme
        if st.button("Seçilenleri Word Belgesi Olarak İndir"):
            doc = Document()
            doc.add_heading("Bireyselleştirilmiş Eğitim Planı", level=1)
            doc.add_paragraph(f"Öğretmen Adı: {teacher}")
            doc.add_paragraph(f"Öğrenci Adı: {student}")
            doc.add_paragraph(f"Grup: {selected_group}")
            doc.add_paragraph(f"Ders: {selected_lesson}")

            if selected_short_terms:
                doc.add_heading("Kısa Vadeli Hedefler", level=2)
                for item in selected_short_terms:
                    doc.add_paragraph(f"- {item}")

            if selected_long_terms:
                doc.add_heading("Uzun Vadeli Hedefler", level=2)
                for item in selected_long_terms:
                    doc.add_paragraph(f"- {item}")

            if selected_teaching_terms:
                doc.add_heading("Öğretimsel Hedefler", level=2)
                for item in selected_teaching_terms:
                    doc.add_paragraph(f"- {item}")

            # 📌 Dosyayı kaydet ve indirme linki oluştur
            file = BytesIO()
            doc.save(file)
            file.seek(0)
            
            st.download_button(label="📥 Word Belgesini İndir", data=file, file_name=f"{student}_BEP.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
